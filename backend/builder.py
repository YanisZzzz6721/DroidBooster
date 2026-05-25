from docx import Document
from docx.shared import Pt, Cm
import copy
import os
 
 
# ─────────────────────────────────────────
# UTILITAIRES
# ─────────────────────────────────────────
 
def _set_text(para, text: str):
    """Remplace le texte en gardant le format du premier run."""
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run.text = ''
 
def _set_bold_colon(para, label: str, desc: str):
    """Écrit 'label : desc' — label en gras, desc normal."""
    size = para.runs[0].font.size if para.runs else None
    # Supprime physiquement tous les anciens runs du XML
    from lxml import etree
    p = para._p
    for r in p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
        p.remove(r)
    r1 = para.add_run(label)
    r1.bold = True
    if size: r1.font.size = size
    r2 = para.add_run(' : ' + desc)
    r2.bold = False
    if size: r2.font.size = size
 
def _remove_para(para):
    """Supprime physiquement le paragraphe du document."""
    p = para._element
    p.getparent().remove(p)
 
 
# ─────────────────────────────────────────
# TITRE
# ─────────────────────────────────────────
 
def _fill_titre(doc, cv):
    for para in doc.paragraphs:
        if '{{Intitulé_poste}}' in para.text or '{Intitulé_poste}' in para.text:
            _set_text(para, cv.get('titre', ''))
            return
 
 
# ─────────────────────────────────────────
# PROFIL
# ─────────────────────────────────────────
 
def _fill_profil(doc, cv):
    profil_text = ''
    for s in cv.get('sections', []):
        if s['type'] == 'profil':
            profil_text = s.get('texte', '')
            break
 
    for para in doc.paragraphs:
        if '{Profil}' in para.text or '{profil}' in para.text:
            _set_text(para, profil_text)
            # Alinéa gauche
            para.paragraph_format.left_indent = Cm(0.5)
            return
 
 
# ─────────────────────────────────────────
# COMPÉTENCES
# ─────────────────────────────────────────
 
def _fill_competences(doc, cv):
    items = []
    for s in cv.get('sections', []):
        if s['type'] == 'competences':
            items = s.get('items', [])
            break
 
    # Trouver tous les paragraphes placeholder
    comp_paras = [
        p for p in doc.paragraphs
        if p.text.strip().startswith('{comp') or p.text.strip().startswith('{compt')
    ]
 
    for i, para in enumerate(comp_paras):
        if i < len(items):
            item = items[i]
            label = item.get('label', '')
            desc  = item.get('description', '') or ', '.join(item.get('bullets', []))
            _set_bold_colon(para, label, desc)
        else:
            # Slot en trop → supprimer le paragraphe
            _remove_para(para)
 
 
# ─────────────────────────────────────────
# EXPÉRIENCES
# ─────────────────────────────────────────
 
def _find_experience_blocs(doc):
    """
    Retourne la liste des blocs d'expérience trouvés dans le template.
    Chaque bloc = {'poste': para, 'date': para, 'bullets': [para, ...], 'extras': [para, ...]}
    'extras' = paragraphes vides entre blocs
    """
    paras  = list(doc.paragraphs)
    blocs  = []
    i      = 0
 
    while i < len(paras):
        t = paras[i].text.strip()
        if t == 'Intitulé du poste – Entreprise,Ville':
            bloc = {'poste': paras[i], 'date': None, 'bullets': [], 'extras': []}
            i += 1
            # Ligne date : | Mois AAAA – MoisAAAA |
            if i < len(paras) and paras[i].text.strip().startswith('|'):
                bloc['date'] = paras[i]
                i += 1
            # Bullets jusqu'à ligne vide ou nouveau bloc ou nouvelle section
            while i < len(paras):
                pt = paras[i].text.strip()
                if pt == 'Intitulé du poste – Entreprise,Ville':
                    break
                if pt in ('FORMATION ET DIPLÔMES', 'FORMATION', 'COMPÉTENCES'):
                    break
                if pt == '':
                    bloc['extras'].append(paras[i])
                    i += 1
                    break
                bloc['bullets'].append(paras[i])
                i += 1
            blocs.append(bloc)
        else:
            i += 1
 
    return blocs
 
 
def _fill_experiences(doc, cv):
    experiences = []
    for s in cv.get('sections', []):
        if s['type'] == 'experiences':
            experiences = s.get('items', [])
            break
 
    blocs = _find_experience_blocs(doc)
 
    for idx, bloc in enumerate(blocs):
        if idx < len(experiences):
            exp = experiences[idx]
 
            # Poste
            poste_text = exp['poste']
            if exp.get('entreprise'):
                poste_text += f"  –  {exp['entreprise']}"
            _set_text(bloc['poste'], poste_text)
 
            # Date
            if bloc['date']:
                _set_text(bloc['date'], exp.get('date', ''))
 
            # Bullets : remplir ou SUPPRIMER les lignes en trop
            bullets = exp.get('bullets', [])
            for bi, bp in enumerate(bloc['bullets']):
                if bi < len(bullets):
                    _set_text(bp, bullets[bi])
                else:
                    # Bullet vide → on supprime physiquement la ligne
                    _remove_para(bp)
 
        else:
            # Bloc d'expérience en trop → supprimer tous ses paragraphes
            _remove_para(bloc['poste'])
            if bloc['date']:
                _remove_para(bloc['date'])
            for bp in bloc['bullets']:
                _remove_para(bp)
            for ep in bloc['extras']:
                _remove_para(ep)
 
 
# ─────────────────────────────────────────
# FONCTION PRINCIPALE
# ─────────────────────────────────────────
 
def build_docx(cv: dict, template_path: str, output_path: str) -> str:
    doc = Document(template_path)
 
    _fill_titre(doc, cv)
    _fill_profil(doc, cv)
    _fill_competences(doc, cv)
    _fill_experiences(doc, cv)
 
    os.makedirs(
        os.path.dirname(output_path) if os.path.dirname(output_path) else '.',
        exist_ok=True
    )
    doc.save(output_path)
    return output_path
 