import sys, os, tempfile
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
 
import pytest
from docx import Document
from docx.shared import Pt
from builder import _set_text, _set_bold_colon, _remove_para, build_docx
from parser import parse_cv
 
 
# ── Helpers de test ──────────────────────────────────────────────────────────
 
def make_doc_with_para(text: str) -> tuple:
    """Crée un doc temporaire avec un paragraphe contenant le texte donné."""
    doc = Document()
    para = doc.add_paragraph()
    para.add_run(text)
    return doc, para
 
 
# ── 1. _set_text ─────────────────────────────────────────────────────────────
 
class TestSetText:
 
    def test_remplace_texte(self):
        doc, para = make_doc_with_para("ancien texte")
        _set_text(para, "nouveau texte")
        assert para.text == "nouveau texte"
 
    def test_efface_runs_supplementaires(self):
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("run 1")
        para.add_run(" run 2")
        _set_text(para, "remplacement")
        assert para.text == "remplacement"
 
    def test_para_sans_runs(self):
        doc = Document()
        para = doc.add_paragraph()
        _set_text(para, "texte")
        assert para.text == "texte"
 
    def test_texte_vide(self):
        doc, para = make_doc_with_para("ancien")
        _set_text(para, "")
        assert para.text == ""
 
 
# ── 2. _set_bold_colon ───────────────────────────────────────────────────────
 
class TestSetBoldColon:
 
    def test_format_label_desc(self):
        doc, para = make_doc_with_para("placeholder")
        _set_bold_colon(para, "Accueil", "Orientation clients")
        assert "Accueil" in para.text
        assert "Orientation clients" in para.text
        assert " : " in para.text
 
    def test_label_bold(self):
        doc, para = make_doc_with_para("placeholder")
        _set_bold_colon(para, "Accueil", "desc")
        assert para.runs[0].bold is True
 
    def test_desc_non_bold(self):
        doc, para = make_doc_with_para("placeholder")
        _set_bold_colon(para, "Accueil", "desc")
        assert para.runs[1].bold is False
 
 
# ── 3. _remove_para ──────────────────────────────────────────────────────────
 
class TestRemovePara:
 
    def test_supprime_paragraphe(self):
        doc = Document()
        p1 = doc.add_paragraph("à garder")
        p2 = doc.add_paragraph("à supprimer")
        _remove_para(p2)
        texts = [p.text for p in doc.paragraphs]
        assert "à supprimer" not in texts
        assert "à garder" in texts
 
    def test_compte_paragraphes(self):
        doc = Document()
        paras = [doc.add_paragraph(f"para {i}") for i in range(3)]
        _remove_para(paras[1])
        assert len([p for p in doc.paragraphs if p.text]) == 2
 
 
# ── 4. BUILD_DOCX — intégration ──────────────────────────────────────────────
 
TEMPLATE_PATH = "/mnt/user-data/uploads/Template_cv.docx"
 
CV_3_EXP = """
# Yanis Zouggagh
**Hôte d'accueil — Relation client**
 
## Profil
Étudiant ingénieur disponible immédiatement.
 
## Compétences
- **Accueil & relation client** : Orientation, diplomatie, tous publics
- **Organisation & rigueur** : Procédures, multi-tâches, autonomie
- **Langues** : Français natif, Anglais C1 Cambridge
- **Outils & divers** : Word, Excel, Permis B
- **Gestion situations** : Calme, réactivité, professionnalisme
- **Présentation & image** : Tenue soignée, ponctualité
 
## Expériences
 
### Agent prestataire — SNCF, Strasbourg
Décembre 2024 – Septembre 2025
- Accueil voyageurs en gare
- Gestion demandes et restitution objets
- Enregistrement selon procédures SNCF
- Interface voyageurs-services SNCF
 
### Vendeur conseil — Banette, Strasbourg
Mai 2024 – Novembre 2024
- Accueil chaleureux et conseil
- Encaissement et gestion file
- Application normes hygiène
 
### Serveur — La Couronne
Janvier 2023 – Septembre 2023
- Accueil et prise de commandes
- Coordination équipe cuisine
"""
 
 
@pytest.mark.skipif(
    not os.path.exists(TEMPLATE_PATH),
    reason="Template DOCX non disponible dans cet environnement"
)
class TestBuildDocx:
 
    def setup_method(self):
        self.tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        self.tmp.close()
        self.output = self.tmp.name
 
    def teardown_method(self):
        if os.path.exists(self.output):
            os.unlink(self.output)
 
    def test_genere_fichier(self):
        cv = parse_cv(CV_3_EXP)
        build_docx(cv, TEMPLATE_PATH, self.output)
        assert os.path.exists(self.output)
        assert os.path.getsize(self.output) > 0
 
    def test_fichier_valide_docx(self):
        cv = parse_cv(CV_3_EXP)
        build_docx(cv, TEMPLATE_PATH, self.output)
        doc = Document(self.output)
        assert doc is not None
 
    def test_titre_present(self):
        cv = parse_cv(CV_3_EXP)
        build_docx(cv, TEMPLATE_PATH, self.output)
        doc = Document(self.output)
        texts = [p.text for p in doc.paragraphs]
        full = " ".join(texts)
        assert "Hôte d'accueil" in full or "accueil" in full.lower()
 
    def test_experiences_presentes(self):
        cv = parse_cv(CV_3_EXP)
        build_docx(cv, TEMPLATE_PATH, self.output)
        doc = Document(self.output)
        texts = " ".join(p.text for p in doc.paragraphs)
        assert "SNCF" in texts
        assert "Banette" in texts
        assert "La Couronne" in texts
 
    def test_4eme_experience_absente(self):
        """Le template a 4 slots, le CV 3 — le 4e bloc doit être supprimé."""
        cv = parse_cv(CV_3_EXP)
        build_docx(cv, TEMPLATE_PATH, self.output)
        doc = Document(self.output)
        texts = " ".join(p.text for p in doc.paragraphs)
        assert "Intitulé du poste" not in texts
 
    def test_competences_presentes(self):
        cv = parse_cv(CV_3_EXP)
        build_docx(cv, TEMPLATE_PATH, self.output)
        doc = Document(self.output)
        texts = " ".join(p.text for p in doc.paragraphs)
        assert "Accueil" in texts
        assert "Langues" in texts
 
    def test_slot_competence_vide_absent(self):
        """6 compétences dans le CV, template a 7 — le 7e slot doit disparaître."""
        cv = parse_cv(CV_3_EXP)
        build_docx(cv, TEMPLATE_PATH, self.output)
        doc = Document(self.output)
        texts = " ".join(p.text for p in doc.paragraphs)
        assert "{comp" not in texts and "{compt" not in texts
 
    def test_cree_dossier_output(self):
        tmp_dir = tempfile.mkdtemp()
        output = os.path.join(tmp_dir, "sous_dossier", "cv.docx")
        cv = parse_cv(CV_3_EXP)
        build_docx(cv, TEMPLATE_PATH, output)
        assert os.path.exists(output)
 